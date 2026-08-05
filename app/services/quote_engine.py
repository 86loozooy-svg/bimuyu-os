import json
from typing import Any

from app.database import db_session


def _calc_item_qty(item: dict, area: float) -> float:
    """Calculate item quantity based on qtyMode: factor / fixed / manual."""
    mode = item.get("qtyMode", "factor")
    if mode == "fixed":
        return float(item.get("fixed", 0) or 0)
    if mode == "manual":
        return float(item.get("manualQty", 0) or 0)
    # default: factor
    factor = float(item.get("factor", 0) or 0)
    return area * factor


def calculate_quote_totals(
    json_detail: dict[str, Any],
    design_fee_pct: float,
    management_fee_pct: float,
    tax_pct: float,
    margin_pct: float,
) -> dict[str, float]:
    """Calculate totals — supports both 'groups' (legacy) and 'spaces' mode."""
    direct_cost = 0.0
    mode = json_detail.get("mode", "groups")

    if mode == "spaces":
        for space in json_detail.get("spaces", []):
            area = float(space.get("area", 0) or 0)
            for item in space.get("items", []):
                qty = _calc_item_qty(item, area)
                price = float(item.get("price", 0) or 0)
                line_total = qty * price
                item["quantity"] = round(qty, 4)
                item["line_total"] = line_total
                direct_cost += line_total
    else:
        for group in json_detail.get("groups", []):
            for item in group.get("items", []):
                qty = float(item.get("quantity", 0) or 0)
                unit_price = float(item.get("unit_price", 0) or 0)
                line_total = qty * unit_price
                item["line_total"] = line_total
                direct_cost += line_total

    design_fee = direct_cost * design_fee_pct / 100
    subtotal = direct_cost + design_fee
    management_fee = subtotal * management_fee_pct / 100
    before_tax = subtotal + management_fee
    tax = before_tax * tax_pct / 100
    before_margin = before_tax + tax
    margin = before_margin * margin_pct / 100
    total = before_margin + margin

    return {
        "direct_cost": round(direct_cost, 2),
        "design_fee": round(design_fee, 2),
        "management_fee": round(management_fee, 2),
        "tax": round(tax, 2),
        "margin": round(margin, 2),
        "total": round(total, 2),
    }


def merge_template_with_quantities(template_structure: dict, quantities: dict[str, float]) -> dict:
    """Merge BOQ template with user-entered quantities keyed by 'group_idx-item_idx'."""
    result = {"groups": []}
    for gi, group in enumerate(template_structure.get("groups", [])):
        new_group = {"name": group["name"], "items": []}
        for ii, item in enumerate(group.get("items", [])):
            key = f"{gi}-{ii}"
            qty = quantities.get(key, 0)
            new_item = dict(item)
            new_item["quantity"] = float(qty or 0)
            new_group["items"].append(new_item)
        result["groups"].append(new_group)
    return result


def get_studio_defaults() -> dict[str, float]:
    with db_session() as conn:
        row = conn.execute("SELECT * FROM studio_profile WHERE id = 1").fetchone()
        if not row:
            return {
                "design_fee_pct": 15,
                "management_fee_pct": 8,
                "tax_pct": 6,
                "margin_pct": 25,
            }
        return {
            "design_fee_pct": row["default_design_fee_pct"],
            "management_fee_pct": row["default_management_fee_pct"],
            "tax_pct": row["default_tax_pct"],
            "margin_pct": row["default_margin_pct"],
        }


def _iter_sections(detail: dict):
    """Yield (section_name, area, items) — works for both 'spaces' and 'groups' mode."""
    mode = detail.get("mode", "groups")
    if mode == "spaces":
        for space in detail.get("spaces", []):
            yield (space.get("name", ""), float(space.get("area", 0) or 0), space.get("items", []))
    else:
        for group in detail.get("groups", []):
            yield (group.get("name", ""), 0, group.get("items", []))


def _item_price(item: dict) -> float:
    return float(item.get("price", item.get("unit_price", 0)) or 0)


def export_quote_pdf(quote: dict, project: dict, studio: dict) -> bytes:
    import os
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Load a Unicode font that supports Chinese
    font_path = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
    if os.path.exists(font_path):
        pdf.add_font("UnicodeFont", "", font_path, uni=True)
        font_family = "UnicodeFont"
    else:
        font_family = "Helvetica"

    pdf.set_font(font_family, size=16)
    pdf.cell(0, 10, studio.get("name", "Studio OS") or "Studio OS", ln=True)
    pdf.set_font(font_family, size=12)
    pdf.cell(0, 8, f"项目：{project.get('name', '')}", ln=True)
    pdf.cell(0, 8, f"报价 v{quote.get('version', 1)}  ·  状态：{quote.get('status', 'draft')}", ln=True)
    pdf.ln(5)

    detail = json.loads(quote["json_detail"]) if isinstance(quote["json_detail"], str) else quote["json_detail"]
    for sec_name, area, items in _iter_sections(detail):
        pdf.set_font(font_family, size=11)
        label = sec_name
        if area:
            label += f"  ({area:.1f}㎡)"
        pdf.cell(0, 8, label, ln=True)
        pdf.set_font(font_family, size=10)
        for item in items:
            qty = item.get("quantity", 0)
            unit_price = _item_price(item)
            line = item.get("line_total", qty * unit_price)
            line_text = (
                f"  {item.get('name', '')}  |  {qty} {item.get('unit', '')} "
                f"× {unit_price:,.2f} = ¥{line:,.2f}"
            )
            pdf.cell(0, 6, line_text, ln=True)
        pdf.ln(2)

    pdf.ln(5)
    pdf.set_font(font_family, size=10)
    pdf.cell(0, 8, f"直接成本：¥{quote.get('direct_cost', 0):,.2f}", ln=True)
    pdf.cell(0, 8, f"设计费 ({quote.get('design_fee_pct')}%)：已含", ln=True)
    pdf.cell(0, 8, f"管理费 ({quote.get('management_fee_pct')}%)：已含", ln=True)
    pdf.cell(0, 8, f"税 ({quote.get('tax_pct')}%)：已含", ln=True)
    pdf.cell(0, 8, f"利润 ({quote.get('margin_pct')}%)：已含", ln=True)
    pdf.set_font(font_family, size=12)
    pdf.cell(0, 10, f"总计：¥{quote.get('total', 0):,.2f}", ln=True)

    return pdf.output()


# ---------------------------------------------------------------------------
# Excel export (openpyxl)
# ---------------------------------------------------------------------------
def export_quote_excel(quote: dict, project: dict, studio: dict) -> bytes:
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "报价单"

    # --- Styles ---
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(start_color="3B332A", end_color="3B332A", fill_type="solid")
    group_fill = PatternFill(start_color="F0EBE3", end_color="F0EBE3", fill_type="solid")
    money_fmt = '#,##0.00'

    def _style_header(cell):
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    def _style_group(cell):
        cell.font = Font(bold=True, size=11)
        cell.fill = group_fill
        cell.border = border

    def _style_cell(cell, money=False, bold=False):
        cell.border = border
        if money:
            cell.number_format = money_fmt
            cell.alignment = Alignment(horizontal="right")
        if bold:
            cell.font = Font(bold=True)

    # --- Title block ---
    ws.merge_cells("A1:E1")
    ws["A1"] = studio.get("name", "Studio OS") or "Studio OS"
    ws["A1"].font = Font(bold=True, size=16)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    ws.merge_cells("A2:E2")
    ws["A2"] = f"项目：{project.get('name', '')}  ·  编号：{project.get('code', '')}"
    ws["A2"].font = Font(size=11)

    ws.merge_cells("A3:E3")
    ws["A3"] = (
        f"报价 v{quote.get('version', 1)}  ·  状态：{quote.get('status', 'draft')}"
    )
    ws["A3"].font = Font(size=10, color="888888")

    ws.merge_cells("A4:E4")
    ws["A4"] = (
        f"客户：{project.get('client_name', '—')}  ·  "
        f"日期：{quote.get('created_at', '')}"
    )
    ws["A4"].font = Font(size=10, color="888888")

    row = 6  # leave a blank row

    # --- BOQ detail ---
    detail = json.loads(quote["json_detail"]) if isinstance(quote["json_detail"], str) else quote["json_detail"]

    for sec_name, area, items in _iter_sections(detail):
        # Section name
        ws.merge_cells(f"A{row}:E{row}")
        label = sec_name
        if area:
            label += f"  ({area:.1f}㎡)"
        ws[f"A{row}"] = label
        _style_group(ws[f"A{row}"])
        row += 1

        # Column headers
        headers = ["项目", "单位", "单价", "数量", "小计"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=header)
            _style_header(cell)
        row += 1

        # Items
        for item in items:
            qty = float(item.get("quantity", 0) or 0)
            unit_price = _item_price(item)
            line_total = item.get("line_total", qty * unit_price)

            ws.cell(row=row, column=1, value=item.get("name", ""))
            ws.cell(row=row, column=2, value=item.get("unit", ""))
            ws.cell(row=row, column=3, value=unit_price)
            ws.cell(row=row, column=4, value=qty)
            ws.cell(row=row, column=5, value=float(line_total))

            _style_cell(ws.cell(row=row, column=1))
            _style_cell(ws.cell(row=row, column=2))
            _style_cell(ws.cell(row=row, column=3), money=True)
            _style_cell(ws.cell(row=row, column=4), money=True)
            _style_cell(ws.cell(row=row, column=5), money=True)
            row += 1

        row += 1  # blank row between sections

    # --- Summary ---
    summary_start = row
    summaries = [
        ("直接成本", float(quote.get("direct_cost", 0) or 0)),
        (f"设计费 ({quote.get('design_fee_pct', 0)}%)", "含"),
        (f"管理费 ({quote.get('management_fee_pct', 0)}%)", "含"),
        (f"税 ({quote.get('tax_pct', 0)}%)", "含"),
        (f"利润 ({quote.get('margin_pct', 0)}%)", "含"),
        ("总计", float(quote.get("total", 0) or 0)),
    ]
    for label, value in summaries:
        ws.merge_cells(f"A{row}:D{row}")
        ws[f"A{row}"] = label
        ws[f"A{row}"].alignment = Alignment(horizontal="right")
        _style_cell(ws[f"A{row}"], bold=(label == "总计"))
        cell = ws.cell(row=row, column=5, value=value)
        if isinstance(value, (int, float)):
            _style_cell(cell, money=True, bold=(label == "总计"))
        else:
            _style_cell(cell, bold=(label == "总计"))
        row += 1

    # --- Column widths ---
    ws.column_dimensions["A"].width = 36
    ws.column_dimensions["B"].width = 10
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 16

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word export (python-docx)
# ---------------------------------------------------------------------------
def export_quote_word(quote: dict, project: dict, studio: dict) -> bytes:
    import io
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Title ---
    title = doc.add_heading(studio.get("name", "Studio OS") or "Studio OS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Subtitle ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"项目：{project.get('name', '')}  ·  {project.get('code', '')}")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"报价 v{quote.get('version', 1)}  ·  状态：{quote.get('status', 'draft')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"客户：{project.get('client_name', '—')}  ·  日期：{quote.get('created_at', '')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # --- BOQ detail ---
    detail = json.loads(quote["json_detail"]) if isinstance(quote["json_detail"], str) else quote["json_detail"]

    for sec_name, area, items in _iter_sections(detail):
        heading = sec_name
        if area:
            heading += f"  ({area:.1f}㎡)"
        doc.add_heading(heading, level=2)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = "项目"
        hdr[1].text = "单位"
        hdr[2].text = "单价"
        hdr[3].text = "数量"
        hdr[4].text = "小计"

        # Items
        for item in items:
            qty = float(item.get("quantity", 0) or 0)
            unit_price = _item_price(item)
            line_total = item.get("line_total", qty * unit_price)

            cells = table.add_row().cells
            cells[0].text = item.get("name", "")
            cells[1].text = item.get("unit", "")
            cells[2].text = f"{unit_price:,.2f}"
            cells[3].text = f"{qty:g}"
            cells[4].text = f"{float(line_total):,.2f}"

        doc.add_paragraph()  # spacer between sections

    # --- Summary ---
    doc.add_heading("费用汇总", level=2)
    summary_lines = [
        ("直接成本", f"¥{float(quote.get('direct_cost', 0) or 0):,.2f}"),
        (f"设计费 ({quote.get('design_fee_pct', 0)}%)", "已含"),
        (f"管理费 ({quote.get('management_fee_pct', 0)}%)", "已含"),
        (f"税 ({quote.get('tax_pct', 0)}%)", "已含"),
        (f"利润 ({quote.get('margin_pct', 0)}%)", "已含"),
    ]
    for label, value in summary_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"{label}：").bold = True
        p.add_run(f"  {value}")

    # Total line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"总计：¥{float(quote.get('total', 0) or 0):,.2f}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3B, 0x33, 0x2A)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
